#! python3
# vqaWriter.py - creates word document with vqa template based on target company, angles, screening questions, and client.

import pyinputplus as pyip
import docx
import re

doc = docx.Document()
# data

# angles dictionary

'''allAngles = [
    "%s, for the following questions, please elaborate in 1-2 sentences.\n\nThe client will use these responses to best assess who they'd like to speak with.",
    "DisplayLogic: Display Angle Question based on Connection Question",
]'''
anglesQuestions = {
    "Formers": [
        "Please describe your ability to discuss %s. Please include information about: \nRelevant responsibilities you held. \nExposure to the competitive landscape.",
        "\nPlease list the title of the individual you reported to at %s"
    ],
    "Customers": [
        "Are you currently or were you recently a customer of %s?",
        "I am a current customer at my current employer",
        "I was a customer recently at ____ (list company name)",
        "\nWhat is/was your role in the selection/evaluation process for products from %s?",
        "Sole decision-maker", 
        "Part of a selection committee", 
        "End-user only",
        "Not involved with %s",
        "\nWhat is/was the approximate annual spend of your organization on %s products? ($$/year)",
        "\nWhich providers that compete with %s have you evaluated or do you currently use?",
    ],
    "Competitors": [
        "Are you currently or were you recently a competitor to %s?",
        "I am a current competitor at my current employer",
        "I was a competitor recently at ____ (list company name)",
        "\nWhat is your role in the manufacture/sale of products that compete with %s?",
        "\nWhat end client/customer base do you work with?",
        "\nWhat is the typical annual revenue you oversee? ($$/year)",
        "\nWho do you consider major competitors in this space?",
        "\nWhat uses cases are these products addressing?",
    ],
    "Partners": [
        "Are you currently or were you recently a partner of %s?",
        "I am a current partner at my current employer",
        "I was a partner recently at ____ (list company name)",
        "\nWhich products from %s do you sell?",
        "\nPlease provide a ballpark annual volume ($$/year) with %s.",
        "\nPlease list the competitors to %s that you work with and your annual volume with each (ballpark $$/year)",
        "\nPlease elaborate in 1-2 sentences on your ability to discuss \
%s in terms of current trends, market share, competitive landscape, and how product offerings compare.",
    ],
}


# client dictionary
clientLOCDict = {
    "Glenview": {
        "Description": "There are a couple additional questions, but before you proceed, we want to ensure that you \
will be allowed to consult prior to you completing the rest of the questionnaire.\n\n For this project on the \
topics listed in the previous question, my client requires a letter of consent from any expert who is not \
CEO/Majority Owner/Self-Employed/Retired/Unemployed at all of their current companies. If obtained, this letter \
can remain on your profile for other similar requirements in future projects!\n\nThe next questions are regarding \
these rules, and will help assess if an LOC is required.",
        "Highest Ranking": "Are you the highest-ranking employee (CEO/Majority Owner/Self-Employed/Retired/Unemployed) \
at your organization(s)?",
        "Highest Ranking Answers": "Yes, I am CEO\n Yes, I am Majority Owner\n Yes, I am Self-Employed\nYes, I am Retired\nYes, I am Unemployed\nNo, I do not hold one of these positions",
        "DisplayLogic": "DisplayLogic: Display LOC question if Highest Ranking Answer is No",
        "LOC": 'Would you be willing to obtain a signature from your CEO, Head of Legal, or Head of HR?\n\n \
If yes, I will email you with the details for the LOC (what you need to have signed and an easy template).\n\n \
If no, please click "no" and you will not be required to finish the questionnaire.',
        "LOC Answers": "Yes, I will obtain a letter.\nNo, I will not obtain a letter.",
        "Skip Logic": "Skip: If No, Skip to Referral\nResult: If No, Mark as Not A Fit",
    },
    "Viking": {
        "Description": "There are a couple additional questions, but before you proceed, we want to ensure that you \
will be allowed to consult prior to you completing the rest of the questionnaire.\n\n For this project on the \
topics listed in the previous question, my client requires a letter of consent from any expert who is not \
C-Level (CEO/COO/CFO/CMO/CIO/etc. - Reporting directly to CEO) at all of their current companies. If obtained, this letter \
can remain on your profile for other similar requirements in future projects!\n\nThe next questions are regarding \
these rules, and will help assess if an LOC is required.",
        "Highest Ranking": "Are you the highest-ranking employee (CEO/COO/CFO/CMO/CIO/etc. - Reporting directly to CEO) \
at your organization(s)?",
        "Highest Ranking Answers": "Yes, I hold a C-Level Title (CEO/COO/CFO/CMO/etc.)\nYes, I am Majority Owner\n \
        Yes, I am Self-Employed\nYes, I am Retired\nYes, I am Unemployed\nNo, I do not hold one of these positions",
        "DisplayLogic": "DisplayLogic: Display LOC question if Highest Ranking Answer is No",
        "LOC": "Would you be willing to obtain a signature from your CEO\COO\CFO, Head of Legal, or Head of HR?\n\n \
If yes, I will email you with the details for the LOC (what you need to have signed and an easy template).\n\n",
        "LOC Answers": "Yes, I will obtain a letter.\nNo, I will not obtain a letter.",
        "DisplayLogic": "DisplayLogic: Display Caviat question if LOC is No",
        "Caviat Question": "Are you still working in the same industry as the subject of this project, or are you a current customer or partner of this company?",
        "Caviat Answers": "Yes, same industry\nYes, current customer\nYes, current partner\nNo, not in the same industry and not a customer or partner",
        "Skip Logic": "Skip: If Yes, Skip to Referral\nResult: If Yes, Mark as Not A Fit",
    },
    "Lone Pine": {
        "Description": "There are a couple additional questions, but before you proceed, we want to ensure that you \
will be allowed to consult prior to you completing the rest of the questionnaire.\n\n For this project on the \
topics listed in the previous question, my client requires a letter of consent from any expert who is not \
CEO/COO/CFO/Majority Owner/Self-Employed/Retired/Unemployed at all of their current companies. If obtained, this letter \
can remain on your profile for other similar requirements in future projects!\n\nThe next questions are regarding \
these rules, and will help assess if an LOC is required.",
        "Highest Ranking": "Are you the highest-ranking employee (CEO/COO/CFO/Majority Owner/Self-Employed/Retired/Unemployed) \
at your organization(s)?",
        "DisplayLogic": "DisplayLogic: Display LOC question if Highest Ranking Answer is No",
        "Highest Ranking Answers": "Yes, my title is CEO/COO/CFO\nYes, I am Majority Owner\n \
        Yes, I am Self-Employed\nYes, I am Retired\nYes, I am Unemployed\nNo, I do not hold one of these positions",
        "LOC": 'Would you be willing to obtain a signature from your CEO/COO/CFO, Head of Legal, or Head of HR?\n\n \
If yes, I will email you with the details for the LOC (what you need to have signed and an easy template).\n\n \
If no, please click "no" and you will not be required to finish the questionnaire.',
        "LOC Answers": "Yes, I will obtain a letter.\nNo, I will not obtain a letter.",
        "Skip Logic": "Skip: If No, Skip to Referral\nResult: If No, Mark as Not A Fit",
    },
    "Jericho": {
        "Description": "There are a couple additional questions, but before you proceed, we want to ensure that you \
will be allowed to consult prior to you completing the rest of the questionnaire.\n\n For this project on the \
topics listed in the previous question, my client requires a letter of consent from any expert who is not \
C-Level (CEO/COO/CFO/etc. - Reporting directly to CEO) at all of their current companies. If obtained, this letter \
can remain on your profile for other similar requirements in future projects!\n\nThe next questions are regarding \
these rules, and will help assess if an LOC is required.",
        "Highest Ranking": "Are you the highest-ranking employee (CEO/COO/CFO/etc. - Reporting directly to CEO) \
at your organization(s)?",
        "Highest Ranking Answers": "Yes\nNo",
        "DisplayLogic": "DisplayLogic: Display LOC question if Highest Ranking Answer is No",
        "LOC": "Would you be willing to obtain a signature from your CEO\COO\CFO, Head of Legal, or Head of HR? \
If yes, I will email you with the details for the LOC (what you need to have signed and an easy template).",
        "LOC Answers": "Yes, I will obtain a letter.\nNo, I will not obtain a letter.",
    },
}


# standard questions
standardQuestions = {
    "Referral": "Would you like to refer anyone for this project?\n\nExperts receive a referral bonus \
of $100 each time someone he/she referred completes a phone consultation within the first year of being referred.\n\n \
    If yes, please include their First and Last Name, and their phone number or LinkedIn profile.",
    "ReferralAnswers": "Yes (please provide their information here)\n\n No\n",
}

uniqueQuestionTemplates = {
    "3/3 Can discuss": ["3/3 Can discuss in extreme detail", "2/3 Can discuss in depth", "1/3 Can discuss broadly", "0/3 can not discuss"],
}
# Gathering Data: Target Company or Industry
targetType = pyip.inputMenu(
    ["One Target Company", "Target Industry", "Multiple Target Companies"],
    lettered=True,
)
if targetType == "One Target Company":
    target = pyip.inputStr("What is the name of the target company? ")
if targetType == "Target Industry":
    target = pyip.inputStr("What is the name of the target industry? ")
if targetType == "Multiple Target Companies":
    targetCompanyQuantity = pyip.inputNum("How many companies? ")
    targetCompanyList = []
    for i in range(targetCompanyQuantity):
        targetCompanyList.append(input("Company %i: " % int(i + 1)))

# Gathering Data: Angles
anglesIncludedQuantity = pyip.inputNum(
    "How many angles? (Formers, Customers, Partners, Competitors) "
)
anglesList = []
for i in range(anglesIncludedQuantity):
    angle = pyip.inputMenu(
        ["Formers", "Customers", "Partners", "Competitors", "Experts"], lettered=True
    )
    anglesList.append(angle)
    print("Current Angles List = ", anglesList)

# Any unusual question formats?
unusualQuestion = pyip.inputMenu(prompt="Will you be having a ranking matrix question?\n", choices=["Yes", "No"], lettered=True)

# Gathering Data: Client
client = pyip.inputMenu(
    ["Glenview", "Viking", "Lone Pine", "Jericho", "Other"], lettered=True
)

# Writing Questions and Writing Word Doc

# Intro Question

if targetType == "One Target Company":
    introSentences = (
        "My client is looking to learn more about the company %s from the perspective of experts who are:\n\n"
        % target
        + f" of {target} \n".join(anglesList)
        + f" of {target}"
        + "\n\nIs this topic a fit?"
    )
if targetType == "Target Industry":
    introSentences = (
        "My client is looking to learn more about the %s industry from the perspective of experts who are: "
        % target
        + f"of the {target} industry\n".join(anglesList)
        + "\n\nIs this topic a fit?"
    )
if targetType == "Multiple Target Companies":
    introSentences = (
        "My client is looking to learn more about the following companies: ",
        ", ".join(targetCompanyList) + ".\n\nIs this topic a fit?",
    )
introSentencesAnswers = ["Yes", "No"]
introSentencesLogic = [
    "Skip: If No, Skip to Referral",
    "Result: If No, Mark as Not A Fit",
]
#I cannot get this function to work, local variable referenced before assignment error, so commenting it out.
questionNumberTracker = 0
'''
def questionNumber():
    questionNumberTracker += 1
    doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)
'''
# Intro Question
questionNumberTracker += 1
doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)

doc.add_paragraph(introSentences)
for answer in introSentencesAnswers:
    doc.add_paragraph(answer)
for logic in introSentencesLogic:
    doc.add_paragraph(logic)

# Connection Question
questionNumberTracker += 1
doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)
doc.add_paragraph(
    "Which of the following options best describes your connection to the topic?"
)
connectionAnswers = anglesList + ["Other (please elaborate) ", "No connection "]
for answer in connectionAnswers[0:-2]:
    doc.add_paragraph(answer[0:-1] + f" of {target}")
connectionAnswersLogic = [
    "Skip: If No, Skip to Referral",
    "Result: If angle, mark as angle. If No, Mark as Not A Fit",
]
for logic in connectionAnswersLogic:
    doc.add_paragraph(logic)

# Highest Ranking Description and Questions
if client != "Other":
    for key,value in clientLOCDict[client].items():
        if key == "Description" or key == "Highest Ranking" or key == "LOC" or key == "Caviat Question":
            questionNumberTracker += 1
            doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)
        doc.add_paragraph(value)

# angle questions
if unusualQuestion == "yes":
    doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)
    doc.add_paragraph("Scale Matrix Question Template")
    questionNumberTracker += 1
    for key, value in uniqueQuestionTemplates.items():
        for item in value:
            doc.add_paragraph(item)

for value in anglesList:
    questionNumberTracker += 1
    doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)
    for item in anglesQuestions[value]:
        if re.compile("\%s").search(item) != None:
            if targetType == "Multiple Target Companies":
                targetForAngleQ = ", ".join(targetCompanyList)
                doc.add_paragraph(item % targetForAngleQ)
            else:
                doc.add_paragraph(item % target)
        else:
            doc.add_paragraph(item)
    
# Standard Questions
for key, value in standardQuestions.items():
    if key == "Referral":
        questionNumberTracker += 1
        doc.add_paragraph("\nQuestion %i: " % questionNumberTracker)
    doc.add_paragraph(value)

doc.save("vqa.docx")
